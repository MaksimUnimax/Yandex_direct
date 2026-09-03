#!/usr/bin/env python3
"""Build the corrected Step19 physical client package from persisted materialized data.

Outputs are DERIVED client views. Canonical analytical authority remains upstream.
The script is intentionally deterministic at the data/content level and is executed in CI.
"""

from __future__ import annotations

import csv
import hashlib
import json
import textwrap
from pathlib import Path

from openpyxl import Workbook, load_workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.worksheet.table import Table as XLTable, TableStyleInfo
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table as PDFTable, TableStyle
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parent
DATA = ROOT / "step19_correction_materialized"
DATE = "2026-09-03"

XLSX = ROOT / "STEP_19_CLIENT_WORKBOOK_CORRECTED.xlsx"
DOCX = ROOT / "STEP_19_CLIENT_REPORT_CORRECTED.docx"
PDF = ROOT / "STEP_19_CLIENT_REPORT_CORRECTED.pdf"
MANIFEST = ROOT / "STEP_19_PHYSICAL_ARTIFACT_MANIFEST_2026-09-03.json"

BLUE = "1F4E78"
LIGHT = "D9EAF7"
YELLOW = "FFF2CC"
HOLD = "FCE4D6"


def tsv(name):
    with (DATA / name).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f, delimiter="\t"))


def csvrows(name):
    with (DATA / name).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def sha256(path):
    h = hashlib.sha256()
    with Path(path).open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


page_model = tsv("STEP_19_02_BUSINESS_AND_PAGE_MODEL.tsv")
semantic = csvrows("STEP_19_03_SEMANTIC_CORE_MATERIALIZED.csv")
ai_cases = tsv("STEP_19_04_SEARCH_VS_AI_GAP_MATRIX.tsv")
actions = tsv("STEP_19_05_PAGE_ACTION_MAP.tsv")
source_obs = tsv("STEP_19_06_SOURCE_COMPETITOR_OBSERVATIONS.tsv")
priority = tsv("STEP_19_07_PRIORITY_ACTION_PLAN.tsv")
calibration = csvrows("STEP_19_EXECUTION_CALIBRATION_BOARD_112.csv")
measurement = csvrows("STEP_19_MEASUREMENT_PROTOCOL.csv")

assert len(page_model) == 15
assert len(semantic) == 2332
assert sum(r["current_assignment_status"] == "SEARCH_REQUIRED" for r in semantic) == 19
assert len(ai_cases) == 8
assert len(actions) == 34
assert len(priority) == 34
assert len(calibration) == 112
assert len(measurement) == 7
assert sum(r["package_kind"] == "EXACT_ACTION" for r in calibration) == 31
assert sum(r["package_kind"] == "INTERNAL_LINK" for r in calibration) == 15
assert sum(r["package_kind"] == "ROUTE_TO_EXISTING" for r in calibration) == 46
assert sum(r["package_kind"] == "HOLD_RECHECK" for r in calibration) == 20


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------
wb = Workbook()
ws = wb.active
ws.title = "README"

header_fill = PatternFill("solid", fgColor=BLUE)
sub_fill = PatternFill("solid", fgColor=LIGHT)
warn_fill = PatternFill("solid", fgColor=YELLOW)
hold_fill = PatternFill("solid", fgColor=HOLD)
white_bold = Font(color="FFFFFF", bold=True)
blue_bold = Font(color=BLUE, bold=True)

ws.merge_cells("A1:H1")
ws["A1"] = "OKNO-MSK — Step19 corrected client workbook"
ws["A1"].font = Font(size=16, bold=True, color="FFFFFF")
ws["A1"].fill = header_fill
ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
ws.row_dimensions[1].height = 28

for cell, value in [("A3","Ключевой показатель"),("B3","Значение"),("D3","Аналитический приоритет"),("E3","Actions")]:
    ws[cell] = value; ws[cell].fill = header_fill; ws[cell].font = white_bold

kpis = [
    ("Материализованные активные фразы", 2332),
    ("Неразрешённые фразы", 19),
    ("Основные клиентские направления", 15),
    ("Search-vs-AI кейсы", 8),
    ("Page actions", 34),
    ("Execution packages", 112),
    ("HOLD packages", 20),
    ("Measurement classes", 7),
]
for i,(k,v) in enumerate(kpis,4):
    ws.cell(i,1,k); ws.cell(i,2,v); ws.cell(i,1).fill=sub_fill; ws.cell(i,1).font=blue_bold

for i,(k,v) in enumerate([("P1_HIGH",12),("P2_MEDIUM",20),("P3_LATER",1),("HOLD",1)],4):
    ws.cell(i,4,k); ws.cell(i,5,v); ws.cell(i,4).fill=sub_fill; ws.cell(i,4).font=blue_bold

ws["D10"]="Тип work package"; ws["E10"]="Количество"
for c in (ws["D10"],ws["E10"]): c.fill=header_fill; c.font=white_bold
for i,(k,v) in enumerate([("EXACT_ACTION",31),("INTERNAL_LINK",15),("ROUTE_TO_EXISTING",46),("HOLD_RECHECK",20)],11):
    ws.cell(i,4,k); ws.cell(i,5,v); ws.cell(i,4).fill=sub_fill; ws.cell(i,4).font=blue_bold

ws.merge_cells("A14:C14")
ws["A14"]="Как пользоваться файлом"; ws["A14"].fill=header_fill; ws["A14"].font=white_bold
usage = [
    "02_Page_Model — 15 основных направлений и current page roles.",
    "03_Semantic_Core — полный материализованный phrase → task → page view; ручной JOIN repo-файлов не нужен.",
    "05_Page_Actions — конкретно что менять и чего не делать без stronger evidence.",
    "07_Priority_Plan — аналитическая важность; P1/P2/P3 не являются готовым sprint/calendar schedule.",
    "Execution_Calibration — 112 точных packages; owner/effort/capacity остаются TO_CALIBRATE до реального ввода.",
    "Measurement — implementation acceptance, baseline, future metrics, observation window и decision rule.",
    "04_Search_vs_AI — только 8 bounded exact-query cases; не sitewide visibility и не longitudinal tracking.",
    "Workbook = DERIVED/MATERIALIZED client view; canonical authority остаётся upstream.",
    "Correction Step19 provider calls = 0; new paid cost = 0 RUB.",
]
for i,text in enumerate(usage,15):
    ws.cell(i,1,i-14); ws.cell(i,1).fill=sub_fill; ws.cell(i,1).font=blue_bold
    ws.merge_cells(start_row=i,start_column=2,end_row=i,end_column=8); ws.cell(i,2,text)

prov = [
    ("Source branch", "roadmap/kwork-productization-2026-08-28"),
    ("Materialization CI", "run 33752050742"),
    ("Canonical semantic sources", "STEP_08_SEARCH_STAGE_SEMANTIC_SET.tsv + STEP_10_FRESH_R1_ASSIGNMENTS_FINAL.tsv + STEP_11_PHRASE_PAGE_MAP.tsv"),
    ("Canonical action sources", "STEP_18_ACTION_REGISTER.tsv + STEP_18_WORK_PACKAGE_REGISTER.json + STEP_18_HOLD_RECHECK_LEDGER.tsv"),
    ("View status", "DERIVED / MATERIALIZED — not a competing canonical authority"),
    ("Correction rule", "CANONICAL SOURCE != MATERIALIZED CLIENT VIEW; UNKNOWN != ABSENT CALIBRATION PROCESS; TRACEABILITY PASS != CLIENT USABILITY PASS"),
]
ws["A25"]="Provenance"; ws["B25"]="Value"
for c in (ws["A25"],ws["B25"]): c.fill=header_fill; c.font=white_bold
for i,(k,v) in enumerate(prov,26):
    ws.cell(i,1,k); ws.cell(i,1).fill=sub_fill; ws.cell(i,1).font=blue_bold
    ws.merge_cells(start_row=i,start_column=2,end_row=i,end_column=8); ws.cell(i,2,v)

for col,width in {"A":28,"B":24,"C":4,"D":30,"E":16,"F":18,"G":18,"H":18}.items(): ws.column_dimensions[col].width=width
for row in ws.iter_rows():
    for cell in row: cell.alignment=Alignment(vertical="top",wrap_text=True)
ws.freeze_panes="A3"

chart=BarChart(); chart.title="Actions по аналитическому приоритету"; chart.style=10
chart.add_data(Reference(ws,min_col=5,min_row=3,max_row=7),titles_from_data=True)
chart.set_categories(Reference(ws,min_col=4,min_row=4,max_row=7)); chart.height=7; chart.width=12
ws.add_chart(chart,"J2")

sheet_data = [
    ("02_Page_Model",page_model),
    ("03_Semantic_Core",semantic),
    ("04_Search_vs_AI",ai_cases),
    ("05_Page_Actions",actions),
    ("06_Source_Obs",source_obs),
    ("07_Priority_Plan",priority),
    ("Execution_Calibration",calibration),
    ("Measurement",measurement),
]
for sidx,(title,rows) in enumerate(sheet_data,1):
    sh=wb.create_sheet(title)
    headers=list(rows[0].keys())
    sh.append(headers)
    for r in rows: sh.append([r.get(h,"") for h in headers])
    for c in sh[1]: c.fill=header_fill; c.font=white_bold; c.alignment=Alignment(vertical="center",wrap_text=True)
    for row in sh.iter_rows(min_row=2):
        for c in row: c.alignment=Alignment(vertical="top",wrap_text=True)
    sh.freeze_panes="A2"; sh.auto_filter.ref=sh.dimensions
    tab=XLTable(displayName=f"Step19_{sidx}",ref=sh.dimensions)
    tab.tableStyleInfo=TableStyleInfo(name="TableStyleMedium2",showRowStripes=True,showFirstColumn=False,showLastColumn=False)
    sh.add_table(tab)
    for col in range(1,len(headers)+1):
        h=headers[col-1]
        base=18
        if any(x in h for x in ["description","reason","what_to_do","implication","boundary","evidence","acceptance","baseline","metric","decision","url","scope","provenance"]): base=42
        if h in {"phrase","user_task","primary_direction","client_action"}: base=34
        sh.column_dimensions[get_column_letter(col)].width=base
    if title=="Execution_Calibration":
        for row in sh.iter_rows(min_row=2):
            if row[3].value=="HOLD":
                for c in row: c.fill=hold_fill
            else:
                for idx in range(9,14): row[idx].fill=warn_fill
    if title=="03_Semantic_Core":
        status_col=headers.index("current_assignment_status")+1
        for row in range(2,sh.max_row+1):
            if sh.cell(row,status_col).value=="SEARCH_REQUIRED":
                for c in sh[row]: c.fill=warn_fill

wb.save(XLSX)
# Independent reopen
rx=load_workbook(XLSX,read_only=True,data_only=False)
assert rx.sheetnames == ["README","02_Page_Model","03_Semantic_Core","04_Search_vs_AI","05_Page_Actions","06_Source_Obs","07_Priority_Plan","Execution_Calibration","Measurement"]
assert rx["03_Semantic_Core"].max_row-1 == 2332
assert rx["Execution_Calibration"].max_row-1 == 112
assert rx["Measurement"].max_row-1 == 7
rx.close()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
doc=Document()
sec=doc.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.7); sec.right_margin=Inches(.7)
for sty in ["Normal","Title","Heading 1","Heading 2"]:
    doc.styles[sty].font.name="Arial"
doc.styles["Normal"].font.size=Pt(10)
doc.styles["Heading 1"].font.color.rgb=RGBColor(31,78,121); doc.styles["Heading 1"].font.size=Pt(16)
doc.styles["Heading 2"].font.color.rgb=RGBColor(31,78,121); doc.styles["Heading 2"].font.size=Pt(13)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("OKNO-MSK"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor(31,78,121)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Семантика, архитектура страниц и план действий\nпод Yandex Search + bounded Search-vs-AI диагностику"); r.bold=True; r.font.size=Pt(17)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Исправленный клиентский отчёт Step19").italic=True

def dtbl(rows):
    t=doc.add_table(rows=1,cols=2); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.rows[0].cells[0].text="Показатель"; t.rows[0].cells[1].text="Результат"
    for c in t.rows[0].cells:
        for rr in c.paragraphs[0].runs: rr.bold=True
    for k,v in rows:
        cells=t.add_row().cells; cells[0].text=str(k); cells[1].text=str(v)
    return t

dtbl([("Сайт","https://okno-msk.ru/"),("Регион","Москва"),("Active semantic/page map","2332 фразы"),("Основные направления",15),("Search-vs-AI cases",8),("Page actions",34),("Execution packages",112),("Correction provider calls","0; cost 0 RUB")])
doc.add_page_break()
doc.add_heading("1. Краткий итог",1)
doc.add_paragraph("Главный вывод: сайту не нужен механический рост числа посадочных страниц. Сильнейший путь — сохранить доказанных владельцев задач, точнее развести роли близких существующих страниц, маршрутизировать узкие задачи на найденные specialist pages и закрыть ограниченное число доказанных content gaps.")
dtbl([("Supported new-page actions",0),("Supported destructive actions",0),("P1_HIGH",12),("P2_MEDIUM",20),("P3_LATER",1),("HOLD", "1 accounting action / 20 exact packages"),("Final sprint/calendar schedule","NOT READY — calibration required")])
for h,items in [("Что делать аналитически в первую очередь",["Закрепить exact specialist owners перед downstream routing/links.","Развести materially overlapping pages без автоматического merge/delete/redirect.","Усилить существующие страницы по конкретным доказанным content needs.","Не снимать HOLD до закрытия named blocker."]),("Что пакет не обещает",["Гарантированный рост позиций/трафика/лидов/выручки.","Sitewide или longitudinal AI visibility по 8 snapshot cases.","Готовый календарь разработки без реальной оценки исполнителя."])]:
    doc.add_heading(h,2)
    for x in items: doc.add_paragraph(x,style="List Bullet")

doc.add_heading("2. Бизнес-модель и архитектура страниц",1)
doc.add_paragraph("Карта сведена к 15 основным направлениям. Broad owners сохраняются для широких задач, exact specialists становятся primary только для доказанной узкой задачи.")
t=doc.add_table(rows=1,cols=4); t.style="Table Grid"
for j,h in enumerate(["ID","Направление","Primary page","Ключевая граница"]): t.rows[0].cells[j].text=h
for r in page_model:
    c=t.add_row().cells
    for j,v in enumerate([r["direction_id"],r["primary_direction"],r["primary_current_page"],r["important_boundary_or_note"]]): c[j].text=v

doc.add_heading("3. Карта действий по страницам",1)
doc.add_paragraph("P1 означает сильнейшую аналитическую важность, а не первый sprint. До production scheduling нужно заполнить owner/effort/capacity в workbook.")
for r in actions:
    if r["priority"]=="P1_HIGH":
        p=doc.add_paragraph(); p.add_run(r["action_id"]+" — ").bold=True; p.add_run(r["client_action"]+" | Не делать: "+r["do_not_do_boundary"])

doc.add_heading("4. Полный semantic/page workbook",1)
doc.add_paragraph("Физический XLSX содержит материализованные 2332 active phrase → task → page rows. Клиенту больше не нужно вручную соединять Step8/10/11. XLSX — DERIVED view; canonical authority остаётся upstream.")
dtbl([("Master upstream keys",2840),("Active materialized rows",2332),("SEARCH_REQUIRED",19),("Silent drops",0)])
doc.add_paragraph("WORDSTAT OBSERVED COUNT != GUARANTEED EXACT QUERY FREQUENCY").alignment=WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("5. Search-vs-AI: 8 bounded diagnostics",1)
dtbl([("CHANGE",0),("DE_RISK",4),("NO_CHANGE",3),("INSUFFICIENT",1),("Architecture changes from AI",0)])
for r in ai_cases:
    p=doc.add_paragraph(); p.add_run(r["case_id"]+" — "+r["exact_query"]+": ").bold=True; p.add_run(r["architecture_verdict"]+"; "+r["content_verdict"]+". "+r["client_implication"])

doc.add_heading("6. Execution calibration: 112 точных пакетов",1)
dtbl([("Exact action",31),("Internal links",15),("Route to existing",46),("HOLD/recheck",20),("TOTAL",112)])
doc.add_paragraph("UNKNOWN не заменяется догадкой. Non-HOLD fields implementation_owner/effort/capacity/observation window остаются TO_CALIBRATE до реального ввода исполнителя/клиента. P1 != первый sprint.")

doc.add_heading("7. Как измерять внедрение",1)
doc.add_paragraph("Measurement sheet разводит implementation acceptance, baseline, future signals, observation window и decision rule. Числовые targets не выдумывались.")
for r in measurement:
    p=doc.add_paragraph(); p.add_run(r["measurement_class"]+" — ").bold=True; p.add_run(r["purpose"]+" "+r["decision_rule"])

doc.add_heading("8. Ограничения и использование",1)
for x in ["Base-public mode: private Webmaster/Metrika/Direct evidence не использовались.","8 AI cases — bounded diagnostic sample, не sitewide выборка.","HOLD — named blocker, не low value/rejection.","Нет ranking/traffic/lead/revenue guarantees."]:
    doc.add_paragraph(x,style="List Bullet")
doc.add_paragraph("Практический порядок: отчёт → 05_Page_Actions → Execution_Calibration → заполнение owner/effort/capacity/dependencies → production sequence → Measurement. Step20 остаётся отдельным финальным QA.")

doc.add_heading("9. Методические ссылки correction",1)
for x in ["https://searchengineland.com/make-seo-reports-more-actionable-479746","https://ahrefs.com/blog/keyword-mapping/","https://ahrefs.com/blog/seo-topical-map/","https://www.semrush.com/blog/what-is-an-seo-report/","https://yandex.ru/support/webmaster/ru/service/popular-queries","https://yandex.ru/support/webmaster/ru/service/queries-export","https://yandex.ru/support/webmaster/ru/recommendations/site-structure"]:
    doc.add_paragraph(x,style="List Bullet")
doc.save(DOCX)
# reopen QA
rd=Document(DOCX); assert len(rd.paragraphs)>30; assert len(rd.tables)>=4


# ---------------------------------------------------------------------------
# PDF — same decision content in compact standalone form
# ---------------------------------------------------------------------------
font="Helvetica"
font_bold="Helvetica-Bold"
for candidate in [Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf")]:
    if candidate.exists():
        pdfmetrics.registerFont(TTFont("ClientSans",str(candidate))); font="ClientSans"; break
for candidate in [Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf")]:
    if candidate.exists():
        pdfmetrics.registerFont(TTFont("ClientSansBold",str(candidate))); font_bold="ClientSansBold"; break
styles=getSampleStyleSheet()
styles.add(ParagraphStyle(name="CTitle",fontName=font_bold,fontSize=22,leading=27,textColor=colors.HexColor("#1F4E78"),alignment=TA_CENTER,spaceAfter=8))
styles.add(ParagraphStyle(name="CSub",fontName=font_bold,fontSize=14,leading=18,textColor=colors.HexColor("#1F4E78"),spaceBefore=8,spaceAfter=5))
styles.add(ParagraphStyle(name="CBody",fontName=font,fontSize=9.2,leading=12.2,spaceAfter=4))
styles.add(ParagraphStyle(name="CBullet",fontName=font,fontSize=9.2,leading=12.2,leftIndent=12,bulletIndent=2,spaceAfter=2))

def P(txt,style="CBody"): return Paragraph(str(txt).replace("&","&amp;"),styles[style])
def pdf_table(rows,widths=None):
    data=[[P(a),P(b)] for a,b in rows]
    t=PDFTable(data,colWidths=widths or [78*mm,90*mm],repeatRows=1)
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1F4E78")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("FONTNAME",(0,0),(-1,-1),font),("GRID",(0,0),(-1,-1),.4,colors.grey),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#EAF3F8")])]))
    return t

story=[P("OKNO-MSK","CTitle"),P("Семантика, архитектура страниц и план действий под Yandex Search + bounded Search-vs-AI диагностику","CSub"),Spacer(1,6*mm)]
story.append(pdf_table([("Показатель","Результат"),("Active semantic/page map","2332"),("Основные направления","15"),("Search-vs-AI cases","8"),("Page actions","34"),("Execution packages","112"),("Correction provider calls","0; cost 0 RUB")]))
story += [PageBreak(),P("1. Краткий итог","CSub"),P("Сайту не нужен механический рост числа посадочных страниц. Основной путь — сохранить доказанных владельцев, развести роли близких страниц, использовать exact specialists и закрыть bounded content gaps."),pdf_table([("Показатель","Результат"),("Supported new-page actions","0"),("Supported destructive actions","0"),("P1/P2/P3/HOLD","12 / 20 / 1 / 1"),("Final schedule","PENDING_CALIBRATION")])]
for x in ["Закрепить specialist owners перед downstream routing/links.","Развести overlap без автоматического merge/delete/redirect.","Усилить существующие страницы по доказанным content needs.","Не снимать HOLD до закрытия named blocker."]:
    story.append(Paragraph("• "+x,styles["CBullet"]))
story += [P("2. Архитектура: 15 направлений","CSub")]
arch=[["ID","Направление","Primary page"]]+[[r["direction_id"],r["primary_direction"],r["primary_current_page"]] for r in page_model]
t=PDFTable([[P(x) for x in row] for row in arch],colWidths=[17*mm,58*mm,93*mm],repeatRows=1)
t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1F4E78")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("GRID",(0,0),(-1,-1),.35,colors.grey),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#EAF3F8")])]))
story.append(t)
story += [P("3. P1 actions","CSub")]
for r in actions:
    if r["priority"]=="P1_HIGH": story.append(P("<b>"+r["action_id"]+"</b> — "+r["client_action"]+". Не делать: "+r["do_not_do_boundary"]))
story += [P("4. Материализованный semantic workbook","CSub"),P("XLSX содержит 2332 active phrase→task→page rows; SEARCH_REQUIRED=19; silent drops=0. Клиенту не нужен manual JOIN. XLSX — derived view; upstream remains canonical."),P("WORDSTAT OBSERVED COUNT != GUARANTEED EXACT QUERY FREQUENCY")]
story += [P("5. Search-vs-AI","CSub"),pdf_table([("Verdict","Count"),("CHANGE","0"),("DE_RISK","4"),("NO_CHANGE","3"),("INSUFFICIENT","1")])]
for r in ai_cases: story.append(P("<b>"+r["case_id"]+"</b> — "+r["exact_query"]+": "+r["architecture_verdict"]+"; "+r["content_verdict"]+"."))
story += [P("6. Execution calibration","CSub"),pdf_table([("Package kind","Count"),("Exact actions","31"),("Internal links","15"),("Route to existing","46"),("HOLD/recheck","20"),("TOTAL","112")]),P("Unknown owner/effort/capacity remain TO_CALIBRATE. P1 != первый sprint. Production sequence появляется только после реальной calibration.")]
story += [P("7. Measurement","CSub")]
for r in measurement: story.append(P("<b>"+r["measurement_class"]+"</b> — "+r["purpose"]+" "+r["decision_rule"]))
story += [P("8. Ограничения и правильное использование","CSub")]
for x in ["Base-public mode: private Webmaster/Metrika/Direct evidence не использовались.","8 AI cases — bounded diagnostic sample, не sitewide visibility.","HOLD — named blocker, не low value/rejection.","Нет ranking/traffic/lead/revenue guarantees.","Step20 остаётся отдельным Final QA."]:
    story.append(Paragraph("• "+x,styles["CBullet"]))
story += [P("9. Методические ссылки correction","CSub")]
for x in ["https://searchengineland.com/make-seo-reports-more-actionable-479746","https://ahrefs.com/blog/keyword-mapping/","https://ahrefs.com/blog/seo-topical-map/","https://www.semrush.com/blog/what-is-an-seo-report/","https://yandex.ru/support/webmaster/ru/service/popular-queries","https://yandex.ru/support/webmaster/ru/service/queries-export","https://yandex.ru/support/webmaster/ru/recommendations/site-structure"]:
    story.append(P(x))
SimpleDocTemplate(str(PDF),pagesize=A4,rightMargin=16*mm,leftMargin=16*mm,topMargin=14*mm,bottomMargin=14*mm).build(story)
assert len(PdfReader(PDF).pages) >= 3

manifest={
    "date":DATE,
    "step":19,
    "status":"PHYSICAL_PACKAGE_BUILT_AND_OPENABILITY_QA_PASS",
    "view_type":"DERIVED_NOT_CANONICAL",
    "source_data":"step19_correction_materialized",
    "logical_accounting":{"directions":15,"semantic_rows":2332,"unresolved":19,"ai_cases":8,"page_actions":34,"execution_packages":112,"measurement_classes":7},
    "non_fabrication":{"new_provider_calls":0,"new_paid_cost_rub":0.0,"committed_schedule":False,"numeric_targets_invented":False},
    "artifacts":{}
}
for p in [XLSX,DOCX,PDF]:
    manifest["artifacts"][p.name]={"bytes":p.stat().st_size,"sha256":sha256(p)}
manifest["openability"]={"xlsx_reopened":True,"docx_reopened":True,"pdf_pages":len(PdfReader(PDF).pages)}
MANIFEST.write_text(json.dumps(manifest,ensure_ascii=False,indent=2)+"\n",encoding="utf-8")
print(json.dumps(manifest,ensure_ascii=False,indent=2))
