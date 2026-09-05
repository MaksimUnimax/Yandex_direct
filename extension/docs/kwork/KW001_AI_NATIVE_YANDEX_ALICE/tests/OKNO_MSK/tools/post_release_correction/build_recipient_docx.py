#!/usr/bin/env python3
"""Build and style recipient DOCX files from corrected Markdown sources."""

from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve()
JOB = HERE.parents[2]
RELEASE = JOB / "OKNO_MSK_RESEARCH_RELEASE_CORRECTED_2026-09-05"
SOURCES = RELEASE / "sources"
EDITABLE = RELEASE / "editable"

BLUE = "2E74B5"
DARK = "1F4D78"
INK = "0B2545"
MUTED = "5D6773"


def set_font(run, name: str, size: float | None = None, color: str | None = None, bold: bool | None = None):
    run.font.name = name
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trPr.append(node)


def set_table_geometry(table, widths: list[int]):
    total = sum(widths)
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    tblLayout = tblPr.first_child_found_in("w:tblLayout")
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        c = OxmlElement("w:gridCol")
        c.set(qn("w:w"), str(w))
        grid.append(c)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            w = widths[min(i, len(widths) - 1)]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margin(cell)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_font(run, "Calibri", 8.5, MUTED)


def style_doc(path: Path, preset: str, running_label: str):
    doc = Document(path)
    def style_named(name: str):
        return next(s for s in doc.styles if s.name == name)
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)

    body_after = 6
    body_line = 1.10 if preset == "standard_business_brief" else 1.25
    heading = {
        "Heading 1": (16, BLUE, 16 if preset == "standard_business_brief" else 18, 8 if preset == "standard_business_brief" else 10),
        "Heading 2": (13, BLUE, 12 if preset == "standard_business_brief" else 14, 6 if preset == "standard_business_brief" else 7),
        "Heading 3": (12, DARK, 8 if preset == "standard_business_brief" else 10, 4 if preset == "standard_business_brief" else 5),
    }
    normal = style_named("Normal")
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(body_after)
    normal.paragraph_format.line_spacing = body_line
    normal.paragraph_format.widow_control = True
    title_style = style_named("Title")
    title_style.font.name = "Calibri"
    title_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title_style.font.size = Pt(25)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor.from_string(INK)
    title_style.paragraph_format.space_before = Pt(20)
    title_style.paragraph_format.space_after = Pt(14)
    for style_name, (size, color, before, after) in heading.items():
        st = style_named(style_name)
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        if style_name in doc.styles:
            st = style_named(style_name)
            st.font.name = "Calibri"
            st.font.size = Pt(11)
            st.paragraph_format.space_after = Pt(4 if preset == "compact_reference_guide" else 8)
            st.paragraph_format.line_spacing = body_line

    # Pandoc's metadata title acts as the editorial/customer-pack title block.
    first = next((p for p in doc.paragraphs if p.text.strip()), None)
    if first:
        first.paragraph_format.space_before = Pt(20)
        first.paragraph_format.space_after = Pt(14)
        first.paragraph_format.keep_with_next = True
        for run in first.runs:
            set_font(run, "Calibri", 25, INK, True)

    for p in doc.paragraphs:
        p.paragraph_format.widow_control = True
        if p.style and p.style.name.startswith("Heading"):
            p.paragraph_format.keep_with_next = True
        for run in p.runs:
            if run.font.name is None:
                set_font(run, "Calibri")

    # Quiet running header/footer.
    hp = sec.header.paragraphs[0]
    hp.text = running_label
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    for run in hp.runs:
        set_font(run, "Calibri", 8.5, MUTED, True)
    fp = sec.footer.paragraphs[0]
    fp.text = "OKNO_MSK · исправленный выпуск · 2026-09-05 · стр. "
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in fp.runs:
        set_font(run, "Calibri", 8.5, MUTED)
    add_page_field(fp)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        n = len(table.columns)
        if n == 2:
            widths = [2700, 6660]
        elif n == 4:
            widths = [1400, 2800, 2460, 2700]
        elif n == 5:
            widths = [900, 1300, 800, 1700, 4660]
        else:
            base = 9360 // n
            widths = [base] * n
            widths[-1] += 9360 - sum(widths)
        set_table_geometry(table, widths)
        set_repeat_header(table.rows[0])
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if ri == 0:
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "F2F4F7" if preset == "standard_business_brief" else "E8EEF5")
                    cell._tc.get_or_add_tcPr().append(shd)
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.0
                    for run in p.runs:
                        set_font(run, "Calibri", 8.5 if n >= 5 else 9.5, INK if ri == 0 else None, ri == 0)

    core = doc.core_properties
    core.title = running_label
    core.subject = "Исправленный получательский документ OKNO_MSK"
    core.author = "OKNO_MSK research rebuild"
    core.last_modified_by = "OKNO_MSK research rebuild"
    core.comments = "Generated from current post-release canonical authorities; no provider calls."
    doc.save(path)


def build(md_path: Path, out_path: Path, preset: str, label: str):
    out_path.parent.mkdir(parents=True, exist_ok=True)
    source = md_path.read_text(encoding="utf-8")
    lines = source.splitlines()
    if lines and lines[0].startswith("# "):
        lines = lines[1:]
    temp_md = out_path.with_suffix(".pandoc.md")
    temp_md.write_text("\n".join(lines) + "\n", encoding="utf-8")
    subprocess.run([
        "/usr/bin/pandoc", str(temp_md), "--from=gfm", "--to=docx", "--toc", "--toc-depth=2",
        "--metadata", f"title={label}", "--metadata", "lang=ru", "--metadata", "toc-title=Содержание",
        "--output", str(out_path)
    ], check=True)
    temp_md.unlink()
    style_doc(out_path, preset, label)


def main():
    build(
        SOURCES / "01_OKNO_MSK_CLIENT_RESEARCH_REPORT_RU_2026-09-05.md",
        EDITABLE / "01_OKNO_MSK_CLIENT_RESEARCH_REPORT_RU_2026-09-05.docx",
        "standard_business_brief",
        "OKNO_MSK · полный отчёт исследования для владельца",
    )
    build(
        SOURCES / "02_OKNO_MSK_SEO_IMPLEMENTATION_GUIDE_RU_2026-09-05.md",
        EDITABLE / "02_OKNO_MSK_SEO_IMPLEMENTATION_GUIDE_RU_2026-09-05.docx",
        "compact_reference_guide",
        "OKNO_MSK · руководство SEO-специалиста",
    )


if __name__ == "__main__":
    main()
